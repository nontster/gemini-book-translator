"""
Convert JSONL translation output to Microsoft Word document
"""

import json
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def convert_jsonl_to_docx(jsonl_file: str, docx_file: str = None, include_original: bool = False) -> str:
    """
    Convert a JSONL translation file to Microsoft Word document.
    
    Args:
        jsonl_file (str): Path to the input JSONL file
        docx_file (str): Path to the output DOCX file (optional, auto-generated if not provided)
        include_original (bool): If True, include original text before each translation
        
    Returns:
        str: Path to the created DOCX file
    """
    # Auto-generate output filename if not provided
    if docx_file is None:
        docx_file = str(Path(jsonl_file).with_suffix('.docx'))
    
    # Create Word document
    doc = Document()
    
    # Create new Style 'TranslationStyle' 
    styles = doc.styles
    try:
        # Check existing of style
        custom_style = styles['TranslationStyle']
    except KeyError:
        # If no style, create new style
        custom_style = styles.add_style('TranslationStyle', WD_STYLE_TYPE.PARAGRAPH)
        custom_style.base_style = styles['Normal']
    
    # Set Font to the new Style
    font = custom_style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(16)
    
    # Read JSONL file
    pages_processed = 0
    pages_skipped = 0
    
    with open(jsonl_file, 'r', encoding='utf-8') as f:
        for line in f:
            if not line.strip():
                continue
                
            try:
                data = json.loads(line)
            except json.JSONDecodeError:
                continue
            
            page_number = data.get('page_number', '?')
            status = data.get('status', 'unknown')
            translated_text = data.get('translated_text', '')
            original_text = data.get('original_text', '')
            
            # Skip failed or empty pages
            if status != 'success' or not translated_text:
                pages_skipped += 1
                continue
            
            # Add page header
            header = doc.add_paragraph()
            header_run = header.add_run(f"--- Page {page_number} ---")
            header_run.bold = True
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add original text if requested
            if include_original and original_text:
                original_header = doc.add_paragraph()
                original_header.add_run("Original:").bold = True
                
                original_para = doc.add_paragraph(original_text)
                original_para.paragraph_format.left_indent = Inches(0.25)
                
                doc.add_paragraph()  # Spacing
            
            # Add translated text
            if translated_text:
                for line in translated_text.split('\n'):
                    clean_line = line.strip()
                    if clean_line:  # ตรวจสอบว่าไม่ใช่บรรทัดว่าง
                        para = doc.add_paragraph(clean_line)
                        para.style = custom_style

            # Add page break between pages (except last)
            doc.add_paragraph()
            
            pages_processed += 1
    
    # Save document
    doc.save(docx_file)
    
    print(f"✅ Conversion complete!")
    print(f"   Pages converted: {pages_processed}")
    print(f"   Pages skipped: {pages_skipped}")
    print(f"   Output: {docx_file}")
    
    return docx_file


def main():
    parser = argparse.ArgumentParser(description='Convert JSONL translation to Word document')
    parser.add_argument('input', help='Input JSONL file')
    parser.add_argument('-o', '--output', help='Output DOCX file (optional)')
    parser.add_argument('--include-original', action='store_true', 
                        help='Include original text before each translation')
    
    args = parser.parse_args()
    
    if not Path(args.input).exists():
        print(f"❌ Error: File not found: {args.input}")
        return
    
    convert_jsonl_to_docx(args.input, args.output, args.include_original)


if __name__ == "__main__":
    main()
